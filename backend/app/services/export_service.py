import io
from html import escape as html_escape

from fastapi.responses import StreamingResponse

from app.models.content import Content
from app.models.listing import Listing

_ALLOWED_FORMATS = {"txt", "html", "docx", "pdf", "pptx", "flyer_pdf"}


class ExportService:
    async def export(
        self,
        content: Content,
        format: str,
        listing: Listing | None = None,
        branding_settings: dict | None = None,
    ) -> StreamingResponse:
        if format not in _ALLOWED_FORMATS:
            raise ValueError(
                f"Unsupported format: {format}. "
                f"Allowed: {', '.join(sorted(_ALLOWED_FORMATS))}"
            )
        if format == "txt":
            return self._export_txt(content)
        elif format == "html":
            return self._export_html(content)
        elif format == "docx":
            return await self._export_docx(content)
        elif format in ("pptx", "flyer_pdf"):
            return await self._export_flyer(content, format, listing, branding_settings)
        else:
            return await self._export_pdf(content)

    def _export_txt(self, content: Content) -> StreamingResponse:
        buffer = io.BytesIO(content.body.encode("utf-8"))
        return StreamingResponse(
            buffer,
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="content-{content.id}.txt"'},
        )

    def _export_html(self, content: Content) -> StreamingResponse:
        safe_title = html_escape(content.content_type)
        safe_body = html_escape(content.body).replace("\n", "<br>\n")
        html = f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>{safe_title}</title></head>
<body>
<div style="max-width: 800px; margin: 0 auto; font-family: Georgia, serif; padding: 2rem;">
{safe_body}
</div>
</body>
</html>"""
        buffer = io.BytesIO(html.encode("utf-8"))
        return StreamingResponse(
            buffer,
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="content-{content.id}.html"'},
        )

    async def _export_docx(self, content: Content) -> StreamingResponse:
        from docx import Document

        doc = Document()
        doc.add_heading(content.content_type.replace("_", " ").title(), level=1)
        for paragraph in content.body.split("\n\n"):
            doc.add_paragraph(paragraph)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="content-{content.id}.docx"'},
        )

    async def _export_pdf(self, content: Content) -> StreamingResponse:
        safe_title = html_escape(content.content_type.replace("_", " ").title())
        safe_body = html_escape(content.body).replace("\n", "<br>\n")
        html = f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><style>
body {{ font-family: Georgia, serif; max-width: 700px; margin: 0 auto; padding: 2rem; }}
h1 {{ color: #1a365d; }}
</style></head>
<body>
<h1>{safe_title}</h1>
<div>{safe_body}</div>
</body>
</html>"""

        from weasyprint import HTML

        pdf_bytes = HTML(string=html).write_pdf()
        buffer = io.BytesIO(pdf_bytes)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="content-{content.id}.pdf"'},
        )

    async def _export_flyer(
        self,
        content: Content,
        format: str,
        listing: Listing | None,
        branding_settings: dict | None,
    ) -> StreamingResponse:
        from app.services.flyer_service import BrandingConfig, FlyerService

        if not listing:
            raise ValueError("Flyer export requires a listing. Content has no associated listing.")

        branding = BrandingConfig.from_settings(branding_settings or {})
        service = FlyerService(branding)

        listing_data = {
            "address_full": listing.address_full or "",
            "price": float(listing.price) if listing.price else None,
            "bedrooms": listing.bedrooms,
            "bathrooms": float(listing.bathrooms) if listing.bathrooms else None,
            "sqft": listing.sqft,
            "lot_sqft": listing.lot_sqft,
            "year_built": listing.year_built,
            "features": listing.features or [],
            "listing_agent_name": listing.listing_agent_name or "",
            "listing_agent_email": listing.listing_agent_email or "",
            "listing_agent_phone": listing.listing_agent_phone or "",
            "property_type": listing.property_type or "",
        }

        if format == "pptx":
            buffer = service.generate_pptx(listing_data, content.body)
            buffer.seek(0)
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                headers={
                    "Content-Disposition": f'attachment; filename="flyer-{content.id}.pptx"'
                },
            )
        else:
            buffer = service.generate_pdf(listing_data, content.body)
            buffer.seek(0)
            return StreamingResponse(
                buffer,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f'attachment; filename="flyer-{content.id}.pdf"'
                },
            )
